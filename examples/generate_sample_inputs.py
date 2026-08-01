"""Generate sample requirement sources to try fitgap end-to-end.

Usage:  python examples/generate_sample_inputs.py [output_dir]

Creates a fictional "Contoso" BRD (.docx), a backlog (.xlsx), and a workshop
transcript (.vtt) in the output directory (default: sample_data/), then prints
the commands to run.
"""

import sys
from pathlib import Path

from docx import Document
from openpyxl import Workbook

out_dir = Path(sys.argv[1] if len(sys.argv) > 1 else "sample_data")
out_dir.mkdir(parents=True, exist_ok=True)

# --- Word BRD ---------------------------------------------------------------
doc = Document()
doc.add_heading("Contoso CRM Replacement — Business Requirements", level=0)
doc.add_heading("1. Introduction", level=1)
doc.add_paragraph("This document captures requirements from the discovery workshops.")
doc.add_heading("2. Sales Requirements", level=1)
doc.add_paragraph(
    "The system must track opportunities through a configurable sales process with stage gates."
)
doc.add_paragraph(
    "Sales managers need dashboards showing pipeline value by region and salesperson.",
    style="List Bullet",
)
doc.add_paragraph(
    "Quotes must be generated as branded PDF documents and emailed to customers.",
    style="List Bullet",
)
doc.add_heading("3. Customer Service Requirements", level=1)
doc.add_paragraph(
    "Agents must see a unified timeline of all customer interactions across channels.",
    style="List Number",
)
doc.add_paragraph(
    "Cases must be automatically routed to queues based on product and severity.",
    style="List Number",
)
doc.add_paragraph(
    "The system shall send an SMS survey to customers when a case is resolved.",
    style="List Number",
)
table = doc.add_table(rows=3, cols=3)
hdr = table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = "ID", "Requirement", "Priority"
r1 = table.rows[1].cells
r1[0].text, r1[1].text, r1[2].text = (
    "BR-101",
    "Field technicians need offline access to work orders on mobile devices.",
    "Must",
)
r2 = table.rows[2].cells
r2[0].text, r2[1].text, r2[2].text = (
    "BR-102",
    "Invoices must sync to the ERP system within 15 minutes.",
    "Must",
)
doc.save(str(out_dir / "contoso_brd.docx"))

# --- Excel backlog ----------------------------------------------------------
wb = Workbook()
ws = wb.active
ws.append(["Ref", "Requirement Description", "Priority (MoSCoW)", "Workstream"])
ws.append(
    [
        "BL-01",
        "Quotes must be generated as branded PDF documents and emailed to customers",
        "Must",
        "Sales",
    ]
)
ws.append(
    [
        "BL-02",
        "Marketing needs to run customer journeys with email and SMS steps.",
        "Should",
        "Marketing/CI-Journeys",
    ]
)
ws.append(
    [
        "BL-03",
        "Only the finance team may see credit-limit fields on the account record.",
        "Must",
        "Security",
    ]
)
wb.save(str(out_dir / "contoso_backlog.xlsx"))

# --- Workshop transcript ----------------------------------------------------
(out_dir / "workshop.vtt").write_text(
    """WEBVTT

1
00:12:05.000 --> 00:12:14.000
<v Operations Lead>Every quote over 50k has to get director sign-off before it leaves the building.</v>

2
00:14:22.000 --> 00:14:31.000
<v Sales Rep>Honestly we lose hours every week rekeying orders into the ERP by hand.</v>

3
00:31:40.000 --> 00:31:52.000
<v Service Manager>When an engineer closes a job on site, the customer should get a survey pretty much straight away.</v>
""",
    encoding="utf-8",
)

print(f"Sample inputs written to {out_dir}/")
print("\nTry the full pipeline (needs ANTHROPIC_API_KEY):")
print(f"  fitgap run {out_dir}/contoso_brd.docx {out_dir}/contoso_backlog.xlsx -t {out_dir}/workshop.vtt")
print("\nOr without an API key (parsing + register skeleton only):")
print(f"  fitgap ingest {out_dir}/contoso_brd.docx {out_dir}/contoso_backlog.xlsx --no-input")
print("  fitgap report")
