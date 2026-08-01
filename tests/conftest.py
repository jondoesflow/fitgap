"""Shared fixtures: sample source documents built with python-docx / openpyxl."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook

# Requirement statements used across fixtures (docx expectations reference these).
DOCX_SECTION_PARA = (
    "The system must track opportunities through a configurable sales process."
)
DOCX_SECTION_BULLET = "Sales managers need dashboards showing pipeline by region."
DOCX_NUMBERED_1 = "Users must be able to convert leads to opportunities."
DOCX_NUMBERED_2 = (
    "The system shall send an email notification when an opportunity exceeds 100k."
)
DOCX_TABLE_ROW_1 = "Agents must see a unified timeline of customer interactions."
DOCX_TABLE_ROW_2 = "Quotes must be generated as branded PDF documents."

XLSX_ROWS = [
    ("R-10", "Customer service agents need to manage cases across email and phone channels.", "Must", "Customer Service"),
    ("R-11", "The solution must support appointment scheduling for field technicians.", "Should", "Field Service"),
    # Near-duplicate of DOCX_TABLE_ROW_2 — exercises cross-source dedupe.
    ("R-12", "Quotes must be generated as branded PDF documents.", "Must", "Sales"),
]


@pytest.fixture(scope="session")
def sample_docx(tmp_path_factory: pytest.TempPathFactory) -> Path:
    path = tmp_path_factory.mktemp("fixtures") / "sample_requirements.docx"
    doc = Document()
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This document describes the requirements gathered during workshops."
    )
    doc.add_heading("2. Sales Requirements", level=1)
    doc.add_paragraph(DOCX_SECTION_PARA)
    doc.add_paragraph(DOCX_SECTION_BULLET, style="List Bullet")
    doc.add_heading("3. General", level=1)
    doc.add_paragraph(DOCX_NUMBERED_1, style="List Number")
    doc.add_paragraph(DOCX_NUMBERED_2, style="List Number")

    table = doc.add_table(rows=3, cols=3)
    header = table.rows[0].cells
    header[0].text = "ID"
    header[1].text = "Requirement"
    header[2].text = "Priority"
    row1 = table.rows[1].cells
    row1[0].text = "R-1"
    row1[1].text = DOCX_TABLE_ROW_1
    row1[2].text = "Must"
    row2 = table.rows[2].cells
    row2[0].text = "R-2"
    row2[1].text = DOCX_TABLE_ROW_2
    row2[2].text = "Should"

    doc.save(str(path))
    return path


@pytest.fixture(scope="session")
def sample_xlsx(tmp_path_factory: pytest.TempPathFactory) -> Path:
    path = tmp_path_factory.mktemp("fixtures") / "sample_backlog.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["ID", "Requirement", "Priority", "Functional Area"])
    for row in XLSX_ROWS:
        sheet.append(list(row))
    workbook.save(str(path))
    return path


@pytest.fixture(scope="session")
def headerless_xlsx(tmp_path_factory: pytest.TempPathFactory) -> Path:
    """A workbook whose headers cannot be auto-detected."""
    path = tmp_path_factory.mktemp("fixtures") / "cryptic.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Col A", "Col B"])
    sheet.append(["x", "The system must do something specific and useful."])
    workbook.save(str(path))
    return path
